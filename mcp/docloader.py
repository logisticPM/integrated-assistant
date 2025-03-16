#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档加载器模块 - 支持多种文档格式的加载和文本提取
"""

import os
import logging
from typing import List, Optional, Dict, Any
from pathlib import Path
import importlib.util

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("docloader")

class BaseLoader:
    """文档加载器基类"""
    
    def __init__(self):
        pass
    
    def load(self, file_path: str) -> str:
        """
        加载文档并提取文本
        
        Args:
            file_path: 文件路径
        
        Returns:
            提取的文本内容
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        return self._read(file_path)
    
    def _read(self, file_path: str) -> str:
        """
        读取文件内容（由子类实现）
        
        Args:
            file_path: 文件路径
        
        Returns:
            文件内容
        """
        raise NotImplementedError("子类必须实现_read方法")

class TextLoader(BaseLoader):
    """文本文件加载器"""
    
    def _read(self, file_path: str) -> str:
        """
        读取文本文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            文件内容
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

class MarkdownLoader(TextLoader):
    """Markdown文件加载器"""
    pass

class PdfLoader(BaseLoader):
    """PDF文件加载器"""
    
    def __init__(self):
        super().__init__()
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查依赖项"""
        if not importlib.util.find_spec("pypdf"):
            logger.warning("未安装pypdf库，无法处理PDF文件。请运行 pip install pypdf")
    
    def _read(self, file_path: str) -> str:
        """
        读取PDF文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            文件内容
        """
        try:
            import pypdf
            
            text = ""
            with open(file_path, 'rb') as f:
                pdf = pypdf.PdfReader(f)
                for page in pdf.pages:
                    text += page.extract_text() + "\n\n"
            
            return text
        except ImportError:
            logger.error("未安装pypdf库，无法处理PDF文件")
            return f"[无法处理PDF文件 {os.path.basename(file_path)}，请安装pypdf库]"
        except Exception as e:
            logger.exception(f"处理PDF文件时出错: {str(e)}")
            return f"[处理PDF文件 {os.path.basename(file_path)} 时出错: {str(e)}]"

class DocxLoader(BaseLoader):
    """Word文件加载器"""
    
    def __init__(self):
        super().__init__()
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查依赖项"""
        if not importlib.util.find_spec("docx"):
            logger.warning("未安装python-docx库，无法处理Word文件。请运行 pip install python-docx")
    
    def _read(self, file_path: str) -> str:
        """
        读取Word文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            文件内容
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + " | "
                    text += row_text.strip(" | ") + "\n"
                text += "\n"
            
            return text
        except ImportError:
            logger.error("未安装python-docx库，无法处理Word文件")
            return f"[无法处理Word文件 {os.path.basename(file_path)}，请安装python-docx库]"
        except Exception as e:
            logger.exception(f"处理Word文件时出错: {str(e)}")
            return f"[处理Word文件 {os.path.basename(file_path)} 时出错: {str(e)}]"

class CsvLoader(TextLoader):
    """CSV文件加载器"""
    pass

class HtmlLoader(TextLoader):
    """HTML文件加载器"""
    
    def _read(self, file_path: str) -> str:
        """
        读取HTML文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            文件内容
        """
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html = f.read()
            
            soup = BeautifulSoup(html, 'html.parser')
            
            # 移除脚本和样式元素
            for script in soup(["script", "style"]):
                script.extract()
            
            # 获取文本
            text = soup.get_text(separator='\n')
            
            # 处理多余的空行
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            text = '\n'.join(lines)
            
            return text
        except ImportError:
            logger.warning("未安装beautifulsoup4库，将以纯文本方式处理HTML文件")
            return super()._read(file_path)
        except Exception as e:
            logger.exception(f"处理HTML文件时出错: {str(e)}")
            return super()._read(file_path)

def get_loader_for_file(file_path: str) -> BaseLoader:
    """
    根据文件类型获取合适的加载器
    
    Args:
        file_path: 文件路径
    
    Returns:
        文档加载器实例
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext in ['.txt']:
        return TextLoader()
    elif file_ext in ['.md', '.markdown']:
        return MarkdownLoader()
    elif file_ext in ['.pdf']:
        return PdfLoader()
    elif file_ext in ['.docx', '.doc']:
        return DocxLoader()
    elif file_ext in ['.csv']:
        return CsvLoader()
    elif file_ext in ['.html', '.htm']:
        return HtmlLoader()
    else:
        logger.warning(f"未知文件类型: {file_ext}，将尝试以文本方式处理")
        return TextLoader()

def load_document(file_path: str) -> str:
    """
    加载文档并提取文本
    
    Args:
        file_path: 文件路径
    
    Returns:
        提取的文本内容
    """
    try:
        loader = get_loader_for_file(file_path)
        text = loader.load(file_path)
        
        # 处理文本，去除多余空白
        text = text.replace('\r\n', '\n')
        text = '\n\n'.join([para.strip() for para in text.split('\n\n') if para.strip()])
        
        logger.info(f"成功加载文档: {os.path.basename(file_path)}")
        return text
    except Exception as e:
        logger.exception(f"加载文档时出错: {str(e)}")
        return f"[加载文档 {os.path.basename(file_path)} 时出错: {str(e)}]"
